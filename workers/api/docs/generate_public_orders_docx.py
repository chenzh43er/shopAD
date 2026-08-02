# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 若主文件名被 Word 占用，可改写此文件后手动替换
OUT = Path(__file__).resolve().parent / "ShopAD-Public-Order-Lookup-API.docx"

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_cn(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        if level == 1:
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return p


def add_para(text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F5F5F5")
    shd.set(qn("w:val"), "clear")
    p._p.get_or_add_pPr().append(shd)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "D6E3F0")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10)
    doc.add_paragraph()
    return table


title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(
    title.add_run("ShopAD 公开订单查询 API 文档"),
    size=20,
    bold=True,
    color=(0x1F, 0x4E, 0x79),
)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(
    meta.add_run("无需登录 · 无需 Token\n版本日期：2026-08-02"),
    size=10,
    color=(0x66, 0x66, 0x66),
)

add_heading_cn("一、概述", 1)
add_para(
    "本接口用于按订单号或手机号查询订单信息，适用于客服工具、落地页查单等场景。"
    "公开接口不返回审核人、内部备注等敏感字段。"
    "后台其余 /api/orders/*（列表、改状态等）仍需员工 Bearer Token。"
)

add_heading_cn("1.1 Base URL", 2)
add_para("开发 / 线上 Worker：")
add_code("https://shopad-api.ubeator.workers.dev")
add_para("本地开发：")
add_code("http://127.0.0.1:8787")

add_heading_cn("二、按订单号查询", 1)
add_heading_cn("2.1 请求", 2)
add_code("GET /api/orders/by-order-no?order_no={订单号}")
add_para("Query 参数：")
add_table(
    ["参数", "位置", "必填", "说明"],
    [
        ["order_no", "Query", "是*", "订单号，精确匹配"],
        ["orderNo", "Query", "是*", "与 order_no 等价，二选一"],
    ],
)
add_para("* 至少提供 order_no 或 orderNo 其中一个。")

add_heading_cn("2.2 成功响应 200", 2)
add_code(
    """{
  "data": {
    "id": "uuid",
    "order_no": "26080212345678",
    "product_id": "uuid",
    "product_name": "商品名",
    "package_name": "套餐名",
    "customer_name": "收件人",
    "customer_phone": "628123456789",
    "shipping_address": "完整地址",
    "shipping_province": "省",
    "shipping_city": "市",
    "shipping_district": "区",
    "shipping_detail": "详细地址",
    "shipping_order_no": "运单号或 null",
    "total_amount": 199000,
    "status": "awaiting_confirm",
    "payment_type": "cod",
    "created_at": "2026-08-02T10:00:00.000Z",
    "updated_at": "2026-08-02T12:00:00.000Z",
    "currency": {
      "id": "uuid",
      "code": "IDR",
      "name": "Indonesian Rupiah",
      "name_zh": "印尼盾",
      "symbol": "Rp",
      "symbol_suffix": false
    }
  }
}"""
)
add_para("说明：currency 在商品未绑定币种时可能为 null。")

add_heading_cn("2.3 错误响应", 2)
add_table(
    ["状态码", "示例 body", "说明"],
    [
        ["400", '{"error":"请提供订单号"}', "缺少订单号"],
        ["404", '{"error":"订单不存在"}', "无匹配订单"],
        ["500", '{"error":"..."}', "服务端错误"],
    ],
)

add_heading_cn("2.4 调用示例", 2)
add_para("cURL：")
add_code(
    'curl "https://shopad-api.ubeator.workers.dev/api/orders/by-order-no?order_no=26080212345678"'
)
add_para("PowerShell：")
add_code(
    'Invoke-RestMethod "https://shopad-api.ubeator.workers.dev/api/orders/by-order-no?order_no=26080212345678"'
)

add_heading_cn("三、按手机号查询（最近一单）", 1)
add_para(
    "仅返回该手机号下按 updated_at 倒序的最近一笔订单。"
    "响应结构与按订单号查询一致（data 为单个对象，非数组）。"
)
add_heading_cn("3.1 请求", 2)
add_code("GET /api/orders/by-phone?phone={手机号}")
add_para("Query 参数：")
add_table(
    ["参数", "位置", "必填", "说明"],
    [
        ["phone", "Query", "是*", "手机号"],
        ["customer_phone", "Query", "是*", "与 phone 等价，二选一"],
    ],
)
add_para("* 至少提供 phone 或 customer_phone 其中一个。")

add_heading_cn("3.2 手机号匹配规则", 2)
add_para("1. 去掉空格、-、()、.、+")
add_para(
    "2. 去掉前导 0，兼容本地号与国际号（如 08123456789 / 628123456789 / +62 812-3456-789）"
)
add_para("3. 使用包含匹配（ILIKE %digits%）")
add_para("4. 取 updated_at 最新的一条")

add_heading_cn("3.3 成功响应 200", 2)
add_code(
    """{
  "data": {
    "id": "uuid",
    "order_no": "26080212345678",
    "product_name": "商品名",
    "customer_name": "收件人",
    "customer_phone": "628123456789",
    "total_amount": 199000,
    "status": "cod_shipped",
    "payment_type": "cod",
    "created_at": "2026-08-02T10:00:00.000Z",
    "updated_at": "2026-08-02T12:00:00.000Z",
    "currency": { "code": "IDR", "symbol": "Rp" }
  }
}"""
)

add_heading_cn("3.4 错误响应", 2)
add_table(
    ["状态码", "示例 body", "说明"],
    [
        ["400", '{"error":"请提供手机号"}', "缺少手机号"],
        ["404", '{"error":"订单不存在"}', "该手机号无订单"],
        ["500", '{"error":"..."}', "服务端错误"],
    ],
)

add_heading_cn("3.5 调用示例", 2)
add_para("cURL：")
add_code(
    'curl "https://shopad-api.ubeator.workers.dev/api/orders/by-phone?phone=628123456789"'
)
add_para("PowerShell：")
add_code(
    'Invoke-RestMethod "https://shopad-api.ubeator.workers.dev/api/orders/by-phone?phone=628123456789"'
)

add_heading_cn("四、返回字段说明", 1)
add_table(
    ["字段", "类型", "说明"],
    [
        ["id", "string (uuid)", "订单 ID"],
        ["order_no", "string", "订单号"],
        ["product_id", "string | null", "商品 ID"],
        ["product_name", "string | null", "商品名快照"],
        ["package_name", "string | null", "套餐名快照"],
        ["customer_name", "string", "收件人姓名"],
        ["customer_phone", "string | null", "收件人手机号"],
        ["shipping_address", "string | null", "完整地址快照"],
        ["shipping_province", "string | null", "省"],
        ["shipping_city", "string | null", "市"],
        ["shipping_district", "string | null", "区"],
        ["shipping_detail", "string | null", "详细地址"],
        ["shipping_order_no", "string | null", "物流运单号"],
        ["total_amount", "number", "订单金额"],
        ["status", "string", "订单状态"],
        ["payment_type", "string | null", "支付类型"],
        ["created_at", "string (ISO 8601)", "创建时间"],
        ["updated_at", "string (ISO 8601)", "最近更新时间"],
        ["currency", "object | null", "币种信息（随商品）"],
    ],
)

add_heading_cn("4.1 status 枚举", 2)
add_table(
    ["值", "含义"],
    [
        ["pending", "待支付"],
        ["paid", "已支付"],
        ["awaiting_review", "待审核"],
        ["awaiting_confirm", "待确认"],
        ["awaiting_shipment", "待发货"],
        ["shipped", "已发货"],
        ["cod_shipped", "已发货（COD）"],
        ["completed", "已完成"],
        ["cod_completed", "已签收"],
        ["cod_refused", "拒绝签收"],
        ["cancelled", "无效订单"],
    ],
)

add_heading_cn("4.2 payment_type 枚举", 2)
add_table(
    ["值", "含义"],
    [
        ["cod", "货到付款"],
        ["non_cod", "非货到付款"],
    ],
)

add_heading_cn("五、鉴权与安全说明", 1)
add_para("1. 这两个接口不需要 Authorization 请求头。")
add_para("2. 后台其余订单管理接口仍需员工 Bearer Token。")
add_para("3. 公开接口不返回审核人、内部备注、寄件人内部信息等敏感列。")
add_para("4. 手机号查询仅返回最近一单；请仅在可信渠道使用。")

add_heading_cn("六、健康检查（可选）", 1)
add_code("GET /api/health")
add_code(
    """{
  "ok": true,
  "service": "shopad-api",
  "ts": "2026-08-02T13:21:08.042Z"
}"""
)

doc.save(OUT)
print(OUT)
